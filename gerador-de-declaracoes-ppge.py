import os
import re
import unicodedata
import tkinter as tk
from tkinter import filedialog
from docx import Document

def remover_acentos(texto):
    """Remove acentos e caracteres especiais de um texto."""
    try:
        return ''.join(
            c for c in unicodedata.normalize('NFD', str(texto))
            if unicodedata.category(c) != 'Mn'
        )
    except Exception as e:
        print(f"Erro ao remover acentos: {e}")
        return str(texto)

def sanitizar_nome(nome):
    """Remove caracteres inválidos para nomes de arquivos."""
    try:
        nome = remover_acentos(nome)
        nome = re.sub(r'[\\/*?:"<>|]', '', nome)
        nome = re.sub(r'\s+', ' ', nome).strip()
        return nome
    except Exception as e:
        print(f"Erro ao sanitizar nome: {e}")
        return "nome_do_aluno"

def validar_cpf(cpf):
    """Validação básica de formato de CPF."""
    try:
        cpf = re.sub(r'[^0-9]', '', str(cpf))
        return len(cpf) == 11
    except Exception as e:
        print(f"Erro ao validar CPF: {e}")
        return False

def formatar_cpf(cpf):
    """Formata CPF para o padrão 000.000.000-00."""
    try:
        cpf = re.sub(r'[^0-9]', '', str(cpf))
        return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"
    except Exception as e:
        print(f"Erro ao formatar CPF: {e}")
        return "000.000.000-00"

def substituir_mantendo_formatacao(paragrafo, antigo, novo):
    """Substitui texto mantendo a formatação original."""
    try:
        if antigo and novo and antigo in paragrafo.text:
            for run in paragrafo.runs:
                if antigo in run.text:
                    run.text = run.text.replace(str(antigo), str(novo))
    except Exception as e:
        print(f"Erro ao substituir texto: {e}")

def ler_alunos_tabela(tabela):
    """Lê os alunos da tabela do Word, lidando com formatação complexa."""
    alunos = []
    try:
        for row in tabela.rows[1:]:  # Pula o cabeçalho
            try:
                celulas = [cell.text.strip() for cell in row.cells]
                if len(celulas) >= 2 and celulas[0] and celulas[1]:
                    matricula = str(celulas[0])
                    nome = ' '.join(str(celulas[1]).split())
                    nome = nome.replace('\n', ' ').replace('\r', '').strip()
                    alunos.append({
                        "matricula": matricula,
                        "nome": nome
                    })
            except Exception as e:
                print(f"Erro ao processar linha da tabela: {e}")
                continue
    except Exception as e:
        print(f"Erro ao ler tabela: {e}")
    return alunos

def main():
    print("=== GERADOR DE DECLARAÇÕES (VERSÃO ESTÁVEL) ===")
    
    try:
        # Configuração inicial
        root = tk.Tk()
        root.withdraw()
        
        # Seleciona arquivos
        modelo_path = filedialog.askopenfilename(
            title="Selecione o MODELO da declaração",
            filetypes=[("Arquivos Word", "*.docx")]
        )
        if not modelo_path:
            print("Operação cancelada - modelo não selecionado.")
            return
        
        relacao_path = filedialog.askopenfilename(
            title="Selecione a RELAÇÃO de alunos",
            filetypes=[("Arquivos Word", "*.docx")]
        )
        if not relacao_path:
            print("Operação cancelada - relação não selecionada.")
            return
        
        # Extrai nome da turma
        try:
            nome_turma = os.path.basename(relacao_path)
            nome_turma = re.sub(r'\.docx$', '', nome_turma, flags=re.IGNORECASE)
            nome_turma = re.sub(r'(rela[çc][aã]o|lista|alunos|turma)', '', nome_turma, flags=re.IGNORECASE)
            nome_turma = nome_turma.strip() or "Turma Sem Nome"
        except Exception as e:
            print(f"Erro ao extrair nome da turma: {e}")
            nome_turma = "Turma Sem Nome"
        
        # Cria pasta de saída
        saida_dir = os.path.join(os.path.dirname(modelo_path), nome_turma)
        os.makedirs(saida_dir, exist_ok=True)
        
        # Lê relação de alunos
        try:
            relacao_doc = Document(relacao_path)
            alunos = []
            
            for tabela in relacao_doc.tables:
                alunos.extend(ler_alunos_tabela(tabela))
            
            if not alunos:
                raise ValueError("Nenhum aluno encontrado na relação. Verifique se o arquivo contém uma tabela com matrículas e nomes.")
            
            print(f"\nTurma: {nome_turma}")
            print(f"Total de alunos encontrados: {len(alunos)}")
            
        except Exception as e:
            print(f"\n❌ Erro ao ler relação de alunos: {str(e)}")
            print("Verifique se o arquivo contém uma tabela com:")
            print("- Matrículas na primeira coluna")
            print("- Nomes dos alunos na segunda coluna")
            return
        
        # Coleta CPFs
        print("\nDigite os CPFs dos alunos (apenas números):")
        for aluno in alunos:
            while True:
                try:
                    cpf = input(f"{aluno['nome']} ({aluno['matricula']}): ").strip()
                    if validar_cpf(cpf):
                        aluno['cpf'] = formatar_cpf(cpf)
                        break
                    print("CPF inválido! Digite 11 dígitos numéricos.")
                except Exception as e:
                    print(f"Erro ao coletar CPF: {e}")
                    continue
        
        # Gera declarações
        contador = 0
        print("\nGerando declarações...")
        
        for aluno in alunos:
            try:
                doc = Document(modelo_path)
                
                # Substituições
                substituicoes = {
                    "xxxxxxxxxxxxxxxxxxxxx": aluno.get('nome', ''),
                    "xxxxxxxxxxxxxx": aluno.get('cpf', ''),
                    "xxxxx": aluno.get('matricula', '')
                }
                
                for paragraph in doc.paragraphs:
                    for antigo, novo in substituicoes.items():
                        substituir_mantendo_formatacao(paragraph, antigo, novo)
                
                # Nome do arquivo
                nome_arquivo = f"{aluno.get('matricula', '')} - {sanitizar_nome(aluno.get('nome', ''))}.docx"
                caminho_saida = os.path.join(saida_dir, nome_arquivo)
                
                # Evita sobrescrever
                if os.path.exists(caminho_saida):
                    base, ext = os.path.splitext(nome_arquivo)
                    i = 1
                    while os.path.exists(os.path.join(saida_dir, f"{base}_{i}{ext}")):
                        i += 1
                    caminho_saida = os.path.join(saida_dir, f"{base}_{i}{ext}")
                
                doc.save(caminho_saida)
                contador += 1
                print(f"✓ {nome_arquivo}")
                
            except Exception as e:
                print(f"✗ Erro ao gerar declaração para {aluno.get('nome', '')}: {str(e)}")
                continue
        
        print(f"\n✅ Concluído! {contador}/{len(alunos)} declarações geradas.")
        print(f"Pasta de saída: {os.path.abspath(saida_dir)}")
        os.startfile(saida_dir)
    
    except Exception as e:
        print(f"\n❌ ERRO CRÍTICO: {str(e)}")
        print("Recomendações:")
        print("1. Verifique se os arquivos selecionados estão no formato correto")
        print("2. Feche todos os arquivos Word antes de executar")
        print("3. Verifique as permissões da pasta")

if __name__ == "__main__":
    main()